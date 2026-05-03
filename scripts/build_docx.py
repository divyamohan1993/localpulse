"""Rewrite CAPSTONE PROJECT REPORT.docx in place with full content.

The university front matter (title page block, the section titles
Acknowledgement / Abstract / Table of Contents / List of Figures /
List of Tables, the eleven chapter headings, the Questions block and
the References heading) is preserved.  Each section is then filled
with the prose held in scripts/content.py.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Inches

import content as C


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "CAPSTONE PROJECT REPORT.docx"


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def add_page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)


def add_paragraphs(doc, paragraphs, *, italic: bool = False, justify: bool = True):
    for text in paragraphs:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)
        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if italic:
            for run in p.runs:
                run.italic = True


def add_heading_1(doc, text):
    h = doc.add_heading(text, level=1)
    return h


def add_heading_2(doc, text):
    h = doc.add_heading(text, level=2)
    return h


def add_heading_3(doc, text):
    h = doc.add_heading(text, level=3)
    return h


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Paragraph")
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("•  ")
        run.bold = True
        p.add_run(item)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_numbered(doc, items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(style="List Paragraph")
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{i}.  ")
        run.bold = True
        p.add_run(item)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_definition_list(doc, mapping):
    for key, val in mapping.items():
        p = doc.add_paragraph()
        run = p.add_run(f"{key}. ")
        run.bold = True
        p.add_run(val)
        p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_table(doc, rows, *, first_row_header: bool = True, widths_in: list[float] | None = None):
    if not rows:
        return None
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for para in cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
            if r_idx == 0 and first_row_header:
                set_cell_shading(cells[c_idx], "1F2A44")
                for para in cells[c_idx].paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if widths_in:
        for row in table.rows:
            for c_idx, w in enumerate(widths_in):
                if c_idx < len(row.cells):
                    row.cells[c_idx].width = Inches(w)
    return table


def add_code_block(doc, code: str, *, caption: str | None = None):
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cp.runs:
            run.italic = True
            run.font.size = Pt(9)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F4F7")
    pPr.append(shd)
    run = p.add_run(code)
    run.font.name = "Consolas"
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), "Consolas")
    run.font.size = Pt(9)


# ---------------------------------------------------------------------------
# rebuild
# ---------------------------------------------------------------------------

def remove_all_paragraphs_and_tables(doc):
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is sectPr:
            continue
        body.remove(child)


def build():
    doc = Document(str(DOCX))
    remove_all_paragraphs_and_tables(doc)

    # ------------------------------------------------------------------
    # Style tweaks (bigger Calibri so it reads cleanly)
    styles = doc.styles
    base = styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(11)
    base.paragraph_format.space_after = Pt(6)

    for level in (1, 2, 3):
        h = styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)
    styles["Heading 1"].font.size = Pt(20)
    styles["Heading 2"].font.size = Pt(15)
    styles["Heading 3"].font.size = Pt(12)

    # ------------------------------------------------------------------
    # 1. TITLE PAGE
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(C.PROJECT["title"])
    r.bold = True
    r.font.size = Pt(26)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI Crisis Management for Small Communities")
    r.italic = True
    r.font.size = Pt(13)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Synopsis submitted for the partial fulfilment of the degree of")
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BACHELOR OF TECHNOLOGY (CSE)")
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Submitted by")
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(C.PROJECT["author"])
    r.bold = True
    r.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Roll No. {C.PROJECT['roll']}")
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(C.PROJECT["degree"])
    r.font.size = Pt(11)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Mentor: {C.PROJECT['mentor']}")
    r.italic = True
    r.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("YOGANANDA SCHOOL OF AI, COMPUTERS AND DATA SCIENCES")
    r.bold = True
    r.font.size = Pt(13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SHOOLINI UNIVERSITY OF BIOTECHNOLOGY AND MANAGEMENT SCIENCES")
    r.bold = True
    r.font.size = Pt(13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SCIENCES SOLAN, H.P., INDIA")
    r.bold = True
    r.font.size = Pt(13)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(C.PROJECT["year"])
    r.bold = True
    r.font.size = Pt(13)

    add_page_break(doc)

    # ------------------------------------------------------------------
    # 2. ACKNOWLEDGEMENT
    add_heading_1(doc, "Acknowledgement")
    add_paragraphs(doc, C.ACKNOWLEDGEMENT)
    add_page_break(doc)

    # ------------------------------------------------------------------
    # 3. ABSTRACT
    add_heading_1(doc, "Abstract")
    add_paragraphs(doc, C.ABSTRACT)
    p = doc.add_paragraph()
    r = p.add_run("Keywords: ")
    r.bold = True
    p.add_run(
        "crisis management, disaster response, NLP summarisation, multilingual voice "
        "assistant, Twilio, Whisper, GPT-4o, Cloud Run, accessibility, "
        "WCAG 2.2, DPDP Act 2023, Aatmanirbhar Bharat."
    )
    add_page_break(doc)

    # ------------------------------------------------------------------
    # 4. TABLE OF CONTENTS
    add_heading_1(doc, "Table of Contents")
    rows = [("Section", "Page")] + list(C.TOC_ENTRIES)
    add_table(doc, rows, widths_in=[5.0, 1.0])
    add_page_break(doc)

    # ------------------------------------------------------------------
    # 5. LIST OF FIGURES
    add_heading_1(doc, "List of Figures")
    rows = [("Label", "Caption", "Page")]
    for fig in C.FIGURES:
        rows.append(fig)
    add_table(doc, rows, widths_in=[0.9, 4.5, 0.7])
    add_page_break(doc)

    # ------------------------------------------------------------------
    # 6. LIST OF TABLES
    add_heading_1(doc, "List of Tables")
    rows = [("Label", "Caption", "Page")]
    for t in C.TABLES_LIST:
        rows.append(t)
    add_table(doc, rows, widths_in=[0.9, 4.5, 0.7])
    add_page_break(doc)

    # ------------------------------------------------------------------
    # 7. INTRODUCTION
    add_heading_1(doc, "1. Introduction and Problem Definition")
    add_heading_2(doc, "1.1 Background")
    add_paragraphs(doc, C.INTRO_BACKGROUND)
    add_heading_2(doc, "1.2 Problem Statement")
    add_paragraphs(doc, C.INTRO_PROBLEM_STATEMENT)
    add_heading_2(doc, "1.3 Objectives")
    add_numbered(doc, C.INTRO_OBJECTIVES)
    add_heading_2(doc, "1.4 Scope")
    add_paragraphs(doc, C.INTRO_SCOPE)
    add_heading_2(doc, "1.5 Target Users")
    add_bullets(doc, C.INTRO_TARGET_USERS)
    add_heading_2(doc, "1.6 Significance")
    add_paragraphs(doc, C.INTRO_SIGNIFICANCE)
    add_page_break(doc)

    # ------------------------------------------------------------------
    # 8. SYSTEM REQUIREMENTS
    add_heading_1(doc, "2. System Requirements")
    add_heading_2(doc, "2.1 Functional Requirements")
    add_bullets(doc, C.REQ_FUNCTIONAL)
    add_heading_2(doc, "2.2 Non-functional Requirements")
    add_definition_list(doc, C.REQ_NONFUNCTIONAL)
    add_heading_2(doc, "2.3 Hardware and Software Requirements")
    add_definition_list(doc, C.REQ_HARDWARE_SOFTWARE)
    add_page_break(doc)

    # ------------------------------------------------------------------
    # 9. SYSTEM ARCHITECTURE
    add_heading_1(doc, "3. System Architecture and Design")
    add_heading_2(doc, "3.1 Architectural Overview")
    add_paragraphs(doc, C.ARCH_OVERVIEW)
    add_heading_2(doc, "3.2 Component Inventory")
    rows = [("Component", "Description")] + list(C.ARCH_COMPONENTS)
    add_table(doc, rows, widths_in=[2.2, 4.0])
    add_heading_2(doc, "3.3 Data Flow")
    add_paragraphs(doc, C.ARCH_DATAFLOW)
    add_heading_2(doc, "3.4 API Design")
    add_paragraphs(doc, C.ARCH_API_DESIGN)
    add_heading_2(doc, "3.5 Data and State Design")
    add_paragraphs(doc, C.ARCH_DATA_DESIGN)
    add_heading_2(doc, "3.6 Security Architecture")
    add_paragraphs(doc, C.ARCH_SECURITY)
    add_heading_3(doc, "Table 4. STRIDE Threat Model and Mitigations")
    rows = [
        ("Threat (STRIDE)", "Vector", "Mitigation"),
        ("Spoofing", "Fake resident reports", "Anonymous signed-cookie session, CAPTCHA gate on spike."),
        ("Tampering", "Man-in-the-middle on transit", "TLS 1.3, HSTS preload, SRI hashes on CDN scripts."),
        ("Repudiation", "Disputed actions", "Structured audit logs with correlation identifier."),
        ("Information Disclosure", "Leak of caller phone numbers", "AES-256-GCM at rest, no PII in logs/URLs."),
        ("Denial of Service", "Flood of fake reports", "Cloud Armor edge limits, per-route token bucket."),
        ("Elevation of Privilege", "Compromised service account", "Least-privilege IAM, no long-lived keys, WIF."),
    ]
    add_table(doc, rows, widths_in=[1.6, 2.1, 2.5])
    add_heading_2(doc, "3.7 Deployment Topology")
    add_paragraphs(doc, C.ARCH_DEPLOYMENT)
    add_page_break(doc)

    # ------------------------------------------------------------------
    # 10. TECHNOLOGY STACK
    add_heading_1(doc, "4. Technology Stack")
    add_paragraphs(doc, [
        "Table 1 lists every technology choice in the LocalPulse stack with "
        "the reason it was selected over alternatives.  The selection is "
        "guided by three constraints: idle cost close to zero, mobile-first "
        "performance on a 3G connection and accessibility at WCAG 2.2 level "
        "AAA."
    ])
    add_table(doc, C.TECH_STACK, widths_in=[1.4, 2.0, 3.0])
    add_page_break(doc)

    # ------------------------------------------------------------------
    # 11. IMPLEMENTATION
    add_heading_1(doc, "5. Implementation")
    add_heading_2(doc, "5.1 Code Organisation")
    add_paragraphs(doc, C.IMPL_CODE_ORG)
    add_heading_2(doc, "5.2 Public API Endpoints")
    rows = [("Endpoint", "Description")] + list(C.IMPL_KEY_APIS)
    add_table(doc, rows, widths_in=[2.4, 3.8])
    add_heading_2(doc, "5.3 Internationalisation")
    add_paragraphs(doc, C.IMPL_I18N)
    add_heading_2(doc, "5.4 Mobile-First Responsive Layout")
    add_paragraphs(doc, C.IMPL_RESPONSIVE)
    add_heading_2(doc, "5.5 Real-time Update Loop")
    add_paragraphs(doc, C.IMPL_REALTIME)
    add_heading_2(doc, "5.6 AI Summarisation Pipeline")
    add_paragraphs(doc, C.IMPL_AI_PIPELINE)
    add_heading_2(doc, "5.7 Voice Bot Flow")
    add_paragraphs(doc, C.IMPL_VOICE_FLOW)
    add_heading_2(doc, "5.8 Selected Code Listings")
    for snip in C.IMPL_CODE_SNIPPETS:
        add_code_block(doc, snip["code"], caption=snip["caption"])
    add_page_break(doc)

    # ------------------------------------------------------------------
    # 12. ALGORITHMS
    add_heading_1(doc, "6. Algorithms and Models")
    add_heading_2(doc, "6.1 Social-media Summarisation Pipeline")
    add_paragraphs(doc, C.ALGO_SUMMARISATION)
    add_heading_2(doc, "6.2 Voice Intent Classification")
    add_paragraphs(doc, C.ALGO_INTENT)
    add_heading_2(doc, "6.3 Trust Score")
    add_paragraphs(doc, C.ALGO_TRUST)
    add_heading_2(doc, "6.4 Language Identification")
    add_paragraphs(doc, C.ALGO_LANGID)
    add_page_break(doc)

    # ------------------------------------------------------------------
    # 13. TESTING
    add_heading_1(doc, "7. Testing")
    add_paragraphs(doc, [
        "Tests run at six layers, each chosen to catch a specific class of "
        "regression as early as possible in the development cycle."
    ])
    rows = [("Layer", "Tooling", "Coverage")] + list(C.TEST_LAYERS)
    add_table(doc, rows, widths_in=[1.4, 1.6, 3.4])
    add_heading_2(doc, "7.1 Sample Test Results")
    add_table(doc, C.TEST_SAMPLE_RESULTS, widths_in=[2.0, 1.2, 1.2, 1.6])
    add_page_break(doc)

    # ------------------------------------------------------------------
    # 14. RESULTS
    add_heading_1(doc, "8. Results and Performance Analysis")
    add_paragraphs(doc, [
        "The minimum lovable product was measured against a set of "
        "engineering targets that map back to the non-functional "
        "requirements in Section 2.2.  Table 3 reports the target value, "
        "the achieved value and a short note on the measurement context."
    ])
    add_table(doc, C.RESULTS_TARGETS, widths_in=[1.9, 1.2, 1.2, 2.1])
    add_heading_2(doc, "8.1 Discussion")
    add_paragraphs(doc, C.RESULTS_DISCUSSION)
    add_page_break(doc)

    # ------------------------------------------------------------------
    # 15. DEPLOYMENT
    add_heading_1(doc, "9. Deployment")
    add_paragraphs(doc, C.DEPLOY_OVERVIEW)
    add_heading_2(doc, "9.1 Dockerfile")
    add_code_block(doc, C.DEPLOY_DOCKERFILE, caption="Listing 4. Multi-stage Dockerfile.")
    add_heading_2(doc, "9.2 Cloud Run Deploy")
    add_code_block(doc, C.DEPLOY_GCLOUD, caption="Listing 5. gcloud run deploy with custom domain mapping.")
    add_heading_2(doc, "9.3 Continuous Integration and Release")
    add_paragraphs(doc, C.DEPLOY_CICD)
    add_page_break(doc)

    # ------------------------------------------------------------------
    # 16. CHALLENGES
    add_heading_1(doc, "10. Challenges and Solutions")
    for ch in C.CHALLENGES:
        add_heading_2(doc, ch["title"])
        p = doc.add_paragraph()
        r = p.add_run("Problem. ")
        r.bold = True
        p.add_run(ch["problem"])
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph()
        r = p.add_run("Solution. ")
        r.bold = True
        p.add_run(ch["solution"])
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph()
        r = p.add_run("Outcome. ")
        r.bold = True
        p.add_run(ch["outcome"])
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_page_break(doc)

    # ------------------------------------------------------------------
    # 17. CONCLUSION
    add_heading_1(doc, "11. Conclusion and Future Scope")
    add_heading_2(doc, "11.1 Conclusion")
    add_paragraphs(doc, C.CONCLUSION)
    add_heading_2(doc, "11.2 Future Scope")
    add_bullets(doc, C.FUTURE_SCOPE)
    add_paragraphs(doc, C.CONCLUSION_CLOSE, italic=True)
    add_page_break(doc)

    # ------------------------------------------------------------------
    # 18. QUESTIONS (VIVA)
    add_heading_1(doc, "Questions and Answers")
    add_paragraphs(doc, [
        "The following ten questions are the standard viva voce set for "
        "this capstone. Each answer is written to be defensible on the "
        "stand."
    ])
    for i, qa in enumerate(C.VIVA_QA, 1):
        add_heading_2(doc, f"Q{i}. {qa['q']}")
        p = doc.add_paragraph()
        r = p.add_run("Answer. ")
        r.bold = True
        p.add_run(qa["a"])
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_page_break(doc)

    # ------------------------------------------------------------------
    # 19. REFERENCES
    add_heading_1(doc, "References")
    for i, ref in enumerate(C.REFERENCES, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"[{i}] ")
        r.bold = True
        p.add_run(ref)
        p.paragraph_format.space_after = Pt(4)

    # ------------------------------------------------------------------
    # save
    doc.save(str(DOCX))
    size = DOCX.stat().st_size
    print(f"Wrote {DOCX} ({size} bytes, {size/1024:.1f} KiB)")


if __name__ == "__main__":
    build()
